"""
Academic .docx Document Generator for Legal Dissertations (Licență)

This module generates professionally formatted Word documents adhering to strict Romanian academic standards:
- Margins: 2.54 cm (1 inch) all around.
- Font: Times New Roman, 12pt.
- Spacing: 1.5 lines.
- Structure: Cover Page, Title Page, TOC, Introduction, Chapters, Conclusions, Bibliography.
"""

import re
import logging
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement  # Still needed for page numbers
from docx.oxml.ns import qn  # Still needed for page numbers
# FORCE RELOAD 2

from typing import Dict, Any, List
import datetime
import io
import base64
logger = logging.getLogger(__name__)
try:
    import matplotlib.pyplot as plt
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False
    logger.warning("Matplotlib not available. Charts will not be generated.")


# Constants for placeholder values
DEFAULT_INSTITUTION = "UNIVERSITATEA [NUME INSTITUȚIE]\nFACULTATEA DE DREPT"
DEFAULT_SUPERVISOR = "Coordonator Științific:\nProf. Univ. Dr. [NUME PROFESOR]"
DEFAULT_AUTHOR = "Absolvent:\n[NUME STUDENT]"
DEFAULT_CITY_YEAR = f"București\n{datetime.datetime.now().year}"

# Global footnote counter to ensure unique sequential IDs
# Reset for each document generation
_footnote_counter = 0

def generate_academic_docx(report: Dict[str, Any], output_path: str) -> None:
    """
    Generate an academic .docx document from a final report structure.
    ROBUST: Handles missing fields with intelligent fallbacks.
    """
    global _footnote_counter
    _footnote_counter = 0  # Reset counter for each document

    try:
        logger.info(f"Starting .docx generation for report: {report.get('title', 'Untitled')[:50]}...")
        logger.info(f"📊 Report structure check - Keys present: {list(report.keys())}")

        doc = Document()
        _apply_academic_styles(doc)

        # 1. Pagina de Gardă (Cover Page)
        _add_cover_page(doc, "LUCRARE DE LICENȚĂ", is_cover=True)

        # 2. Prima Pagină (Title Page)
        real_title = report.get('title', 'TITLUL LUCRĂRII').upper()
        if not real_title or real_title == 'TITLUL LUCRĂRII':
            logger.warning("⚠️ No title found in report, using fallback")
            real_title = 'ANALIZĂ JURIDICĂ'
        _add_cover_page(doc, real_title, is_cover=False)

        # 3. Cuprins (Table of Contents)
        try:
            _add_table_of_contents(doc, report)
        except Exception as e:
            logger.error(f"❌ Error adding TOC: {e}", exc_info=True)
            # Add minimal TOC
            doc.add_heading('CUPRINS', level=1)
            doc.add_paragraph("[Cuprins generat automat]")
            doc.add_page_break()

        # 4. Introducere
        try:
            if report.get('introduction') and isinstance(report['introduction'], dict):
                _add_introduction(doc, report['introduction'])
            elif report.get('executive_summary'):
                logger.info("Using executive_summary for introduction")
                doc.add_heading('INTRODUCERE', level=1)
                _process_text(doc, report['executive_summary'])
            else:
                logger.warning("⚠️ No introduction found, adding minimal one")
                doc.add_heading('INTRODUCERE', level=1)
                doc.add_paragraph("Prezenta lucrare reprezintă o analiză juridică.")
        except Exception as e:
            logger.error(f"❌ Error adding introduction: {e}", exc_info=True)
            doc.add_heading('INTRODUCERE', level=1)
            doc.add_paragraph("[Introducere - eroare la procesare]")

        # 5. Conținut (Chapters)
        try:
            if report.get('chapters') and isinstance(report['chapters'], list) and len(report['chapters']) > 0:
                _add_chapters(doc, report['chapters'])
            else:
                logger.warning("⚠️ No chapters found, adding placeholder")
                doc.add_heading('CAPITOLUL 1. ANALIZĂ', level=1)
                doc.add_paragraph("Conținut indisponibil.")
        except Exception as e:
            logger.error(f"❌ Error adding chapters: {e}", exc_info=True)
            doc.add_heading('CAPITOLUL 1. CONȚINUT', level=1)
            doc.add_paragraph("[Eroare la procesarea capitolelor]")

        # 5.1. Analiză Vizuală și Comparativă (Tasks)
        if report.get('tasks') and isinstance(report['tasks'], list):
            try:
                _add_tasks_section(doc, report['tasks'])
            except Exception as e:
                logger.error(f"❌ Error adding tasks section: {e}", exc_info=True)

        # 6. Concluzii
        try:
            if report.get('conclusions') and isinstance(report['conclusions'], dict):
                _add_conclusions(doc, report['conclusions'])
            else:
                logger.warning("⚠️ No conclusions found, adding minimal one")
                doc.add_heading('CONCLUZII', level=1)
                doc.add_paragraph("Analiza a fost finalizată.")
        except Exception as e:
            logger.error(f"❌ Error adding conclusions: {e}", exc_info=True)
            doc.add_heading('CONCLUZII', level=1)
            doc.add_paragraph("[Concluzii - eroare la procesare]")

        # 7. Bibliografie
        try:
            biblio = report.get('bibliography', {})
            if not isinstance(biblio, dict):
                logger.warning(f"⚠️ Bibliography is not a dict: {type(biblio)}, creating empty")
                biblio = {}
            _add_bibliography(doc, biblio)
        except Exception as e:
            logger.error(f"❌ Error adding bibliography: {e}", exc_info=True)
            doc.add_page_break()
            doc.add_heading('BIBLIOGRAFIE', level=1)
            doc.add_paragraph("[Bibliografie - eroare la procesare]")

        # Final verification before save
        logger.info("[DOCUMENT SAVE] Preparing to save document...")
        logger.info(f"[DOCUMENT SAVE] Document has {len(doc.part.package.parts)} parts total")

        # Check if footnotes part exists
        try:
            footnote_rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
            footnotes_part = doc.part.part_related_by(footnote_rel_type)
            logger.info(f"[DOCUMENT SAVE] ✅ Footnotes part confirmed: {footnotes_part.partname}")
        except:
            logger.warning("[DOCUMENT SAVE] ⚠️ No footnotes part found in document before save!")

        doc.save(output_path)
        logger.info(f"✅ Successfully generated .docx document at: {output_path}")

    except Exception as e:
        logger.error(f"❌ CRITICAL ERROR generating .docx document: {e}", exc_info=True)
        raise

def _add_tasks_section(doc: Document, tasks: List[Dict[str, Any]]) -> None:
    """
    Adds a section for Chart and Comparison tasks.
    """
    if not tasks: return

    doc.add_heading('ANALIZĂ VIZUALĂ ȘI COMPARATIVĂ', level=1)

    for task in tasks:
        task_type = task.get('type')
        logger.info(f"Processing task: {task.get('id')} ({task_type})")

        # Add some vertical spacing
        doc.add_paragraph()

        if task_type == 'chart':
            if MATPLOTLIB_AVAILABLE:
                _add_chart_task(doc, task)
            else:
                doc.add_paragraph("[Chart generation unavailable - Matplotlib missing]")

        elif task_type == 'comparison':
            _add_comparison_task(doc, task)

    # Page break after this section before Conclusions
    # (Optional, but acts like a chapter end)
    # doc.add_page_break()


def _apply_academic_styles(doc: Document) -> None:
    """
    Configure document-wide academic styles.
    """
    # Margins 2.54 cm
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.page_height = Cm(29.7) # A4
        section.page_width = Cm(21.0)

        # Page numbers in footer (Bottom Center)
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number_field(p)

    # Normal Style: Times New Roman 12, Justified, 1.5 line spacing
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.27)

    # Headings
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_font = h_style.font
        h_font.name = 'Times New Roman'
        h_font.color.rgb = RGBColor(0, 0, 0)
        h_font.bold = True
        pf = h_style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)

        if i == 1:
            h_font.size = Pt(14)
            h_font.all_caps = True
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            h_font.size = Pt(12)
            h_font.all_caps = False
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Block Text for quotes
    if 'Block Text' not in doc.styles:
        bt_style = doc.styles.add_style('Block Text', WD_STYLE_TYPE.PARAGRAPH)
    else:
        bt_style = doc.styles['Block Text']

    bt_font = bt_style.font
    bt_font.name = 'Times New Roman'
    bt_font.size = Pt(11)
    bt_pf = bt_style.paragraph_format
    bt_pf.left_indent = Cm(1.27)
    bt_pf.right_indent = Cm(1.27)
    bt_pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    bt_pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    _ensure_footnote_styles(doc)


def _ensure_footnote_styles(doc: Document) -> None:
    """
    Explicitly define Footnote Text and Footnote Reference styles
    to guarantee academic formatting (10pt, Times New Roman, Indented).
    """
    # 1. Footnote Text (Paragraph Style)
    if 'Footnote Text' not in doc.styles:
        style = doc.styles.add_style('Footnote Text', WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles['Footnote Text']

    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # Academic standard: hanging indent or first line indent?
    # Usually first line indent 1.27cm OR hanging.
    # User requested "Prima linie indentată" (First line indented).
    pf.first_line_indent = Cm(0.5)
    pf.left_indent = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 2. Footnote Reference (Character Style)
    if 'Footnote Reference' not in doc.styles:
        # Character style types are different, often need to access xml directly if add_style restricts
        try:
           ref_style = doc.styles.add_style('Footnote Reference', WD_STYLE_TYPE.CHARACTER)
        except:
           # If it exists but failed check?
           ref_style = doc.styles['Footnote Reference']
    else:
        ref_style = doc.styles['Footnote Reference']

    ref_font = ref_style.font
    ref_font.name = 'Times New Roman'
    ref_font.superscript = True
    # usually 10pt or 12pt superscript, let's keep it auto/12pt base


def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)


def _add_cover_page(doc: Document, center_text: str, is_cover: bool = True) -> None:
    """
    Creates the strict layout:
    - Uppercase Institution (Centered, 12)
    - [Space]
    - CENTER TITLE (16, Caps, Bold) -> 'Lucrare de Licenta' or Actual Title
    - [Space]
    - Supervisor (Left, 12)
    - Student (Right, 12)
    - [Bottom]
    - City/Year (Centered, 12)
    """
    # Header: Institution
    p = doc.add_paragraph(DEFAULT_INSTITUTION)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in p.runs:
        run.font.bold = True
        run.font.size = Pt(12)

    # Spacing (adjust visually based on A4)
    for _ in range(8):
        doc.add_paragraph()

    # Center Text (Type or Title)
    p = doc.add_paragraph(center_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    for run in p.runs:
        run.font.bold = True
        run.font.size = Pt(16)
        run.font.all_caps = True

    # Spacing
    for _ in range(6):
        doc.add_paragraph()

    # Names Table (Supervisor Left, Student Right)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    # Left Cell: Supervisor
    cell_left = table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.text = DEFAULT_SUPERVISOR
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p_left.runs:
        run.font.size = Pt(12)
        run.font.bold = True

    # Right Cell: Student
    cell_right = table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.text = DEFAULT_AUTHOR
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p_right.runs:
        run.font.size = Pt(12)
        run.font.bold = True

    # Push to bottom
    for _ in range(8):
        doc.add_paragraph()

    # City/Year
    p = doc.add_paragraph(DEFAULT_CITY_YEAR)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.bold = False

    doc.add_page_break()


def _add_table_of_contents(doc: Document, report: Dict[str, Any]) -> None:
    doc.add_heading('CUPRINS', level=1)

    # Simple static generation
    _add_toc_entry(doc, "Introducere")

    chapters = report.get('chapters', [])
    for ch in chapters:
        num = ch.get('chapter_number', '')
        title = ch.get('chapter_title', '')
        _add_toc_entry(doc, f"{num}. {title}", level=0)

        for sub in ch.get('subsections', []):
            s_num = sub.get('number', '')
            s_title = sub.get('title', '')
            _add_toc_entry(doc, f"{s_num} {s_title}", level=1)

    _add_toc_entry(doc, "Concluzii")
    _add_toc_entry(doc, "Bibliografie")

    doc.add_page_break()


def _add_toc_entry(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=WD_TAB_LEADER.DOTS)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if level == 0:
        run = p.add_run(text)
        run.font.bold = True
    else:
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(text)


def _add_introduction(doc: Document, intro: Dict[str, Any]) -> None:
    """Add introduction with robust field handling."""
    doc.add_heading('INTRODUCERE', level=1)

    if not intro or not isinstance(intro, dict):
        logger.warning("⚠️ Introduction is empty or not a dict")
        doc.add_paragraph("Introducere indisponibilă.")
        return

    content_parts = []

    # Safely extract all possible fields
    if intro.get('motivation') and isinstance(intro['motivation'], str):
        content_parts.append(intro['motivation'])
    if intro.get('context') and isinstance(intro['context'], str):
        content_parts.append(intro['context'])
    if intro.get('scope') and isinstance(intro['scope'], str):
        content_parts.append(intro['scope'])
    if intro.get('methodology') and isinstance(intro['methodology'], str):
        content_parts.append(intro['methodology'])
    if intro.get('summary') and isinstance(intro['summary'], str) and not content_parts:
        content_parts.append(intro['summary'])

    full_text = "\n\n".join(content_parts)

    if not full_text:
        logger.warning("⚠️ Introduction has no text content, adding placeholder")
        full_text = "Prezentul studiu reprezintă o analiză juridică."

    _process_text(doc, full_text)


def _add_chapters(doc: Document, chapters: List[Dict[str, Any]]) -> None:
    for ch in chapters:
        num = ch.get('chapter_number', '')
        title = ch.get('chapter_title', '')
        doc.add_heading(f"CAPITOLUL {num}. {title.upper()}", level=1)

        if ch.get('content'):
            _process_text(doc, ch['content'])

        for sub in ch.get('subsections', []):
            s_num = sub.get('number', '')
            s_title = sub.get('title', '')
            doc.add_heading(f"{s_num} {s_title}", level=2)
            if sub.get('content'):
                _process_text(doc, sub['content'])


def _add_conclusions(doc: Document, concl: Dict[str, Any]) -> None:
    """Add conclusions with robust field handling."""
    doc.add_heading('CONCLUZII', level=1)

    if not concl or not isinstance(concl, dict):
        logger.warning("⚠️ Conclusions is empty or not a dict")
        doc.add_paragraph("Concluzii indisponibile.")
        return

    text = ""
    if concl.get('summary_findings') and isinstance(concl['summary_findings'], str):
        text += concl['summary_findings'] + "\n\n"
    if concl.get('final_perspective') and isinstance(concl['final_perspective'], str):
        text += concl['final_perspective']
    if not text and concl.get('summary') and isinstance(concl['summary'], str):
        text = concl['summary']

    if not text:
        logger.warning("⚠️ Conclusions has no text content, adding placeholder")
        text = "Analiza a fost finalizată cu succes."

    _process_text(doc, text)


def _add_bibliography(doc: Document, biblio: Dict[str, Any]) -> None:
    """Add bibliography with robust handling of various formats (dict/int/str)."""
    doc.add_page_break()
    doc.add_heading('BIBLIOGRAFIE', level=1)

    if not biblio or not isinstance(biblio, dict):
        logger.warning("⚠️ Bibliography is empty or not a dict")
        doc.add_paragraph("Nu există surse citate.")
        return

    jurisprudence = biblio.get('jurisprudence', [])
    if not jurisprudence or not isinstance(jurisprudence, list):
        logger.warning("⚠️ No jurisprudence in bibliography")
        doc.add_paragraph("Nu există surse citate.")
        return

    doc.add_heading('I. Jurisprudență', level=2)

    # Process items - handle dict/int/str formats
    valid_items = []
    for item in jurisprudence:
        if isinstance(item, dict) and item.get('citation'):
            valid_items.append(item)
        elif isinstance(item, (int, str)):
            # Format: just ID - create minimal citation
            valid_items.append({'citation': f"Speță #{item}"})
        else:
            logger.warning(f"⚠️ Skipping invalid bibliography item: {type(item)}")

    if not valid_items:
        logger.warning("⚠️ No valid citations found after processing")
        doc.add_paragraph("Nu există surse valide citate.")
        return

    # Sort by citation text
    try:
        valid_items.sort(key=lambda x: x.get('citation', ''))
    except Exception as e:
        logger.warning(f"⚠️ Could not sort bibliography: {e}")

    for idx, item in enumerate(valid_items, 1):
        citation = item.get('citation', '')
        if citation:
            p = doc.add_paragraph(f"{idx}. {citation}")
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-1.0)

    logger.info(f"✅ Added {len(valid_items)} bibliography entries")


def _process_text(doc: Document, text: str) -> None:
    if not text: return
    paragraphs = text.split('\n\n')
    for p_text in paragraphs:
        p_text = p_text.strip()
        if not p_text: continue

        if p_text.startswith('>') or (len(p_text) > 300 and ('"' in p_text or '„' in p_text)):
            style = 'Block Text'
            p_text = p_text.replace('>', '').strip()
        else:
            style = 'Normal'

        p = doc.add_paragraph(style=style)
        _process_inline_citations(p, p_text)


def _process_inline_citations(paragraph, text: str) -> None:
    """
    Process [[CITATION:ID:Title]] markers and inject footnotes using FULL MANUAL XML.
    Uses global footnote counter to ensure unique sequential IDs.
    """
    pattern = r'\[\[CITATION:(\d+):(.*?)(?:\]\])'
    matches = list(re.finditer(pattern, text))

    if matches:
        logger.info(f"Found {len(matches)} citations in text segment: {text[:50]}...")

    last_idx = 0
    for match in matches:
        # Add text before citation
        pre_text = text[last_idx:match.start()]
        if pre_text:
            paragraph.add_run(pre_text)

        # Extract citation details
        citation_id = match.group(1)
        citation_text = match.group(2)  # e.g., "Cited Case #1001"

        try:
            # FULL MANUAL IMPLEMENTATION with global counter
            _inject_manual_footnote(paragraph, citation_text, citation_id)
        except Exception as e:
            logger.error(f"❌ [MANUAL ERROR] Failed for citation {citation_id}: {e}", exc_info=True)
            # Fallback: simple superscript
            run = paragraph.add_run(f"[{citation_id}]")
            run.font.superscript = True

        last_idx = match.end()

    # Add remaining text after last citation
    if last_idx < len(text):
        paragraph.add_run(text[last_idx:])


def _inject_manual_footnote(paragraph, text: str, citation_id: str) -> None:
    """
    Manually inject footnote using XML with global counter for unique IDs.
    This ensures sequential footnote numbering: 1, 2, 3, 4... (no duplicates!)
    """
    global _footnote_counter

    # Get document part from paragraph
    part = paragraph.part
    if not hasattr(part, 'package'):
        logger.error("[MANUAL] Paragraph has no package")
        return

    # Get/create footnotes part
    from docx.opc.part import XmlPart
    from docx.opc.packuri import PackURI
    from docx.oxml import parse_xml

    footnote_rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"

    try:
        footnotes_part = part.package.main_document_part.part_related_by(footnote_rel_type)
    except KeyError:
        # Create footnotes part
        package = part.package
        partname = PackURI("/word/footnotes.xml")

        xml_content = b"""<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:type="separator" w:id="-1">
    <w:p>
      <w:r>
        <w:separator/>
      </w:r>
    </w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p>
      <w:r>
        <w:continuationSeparator/>
      </w:r>
    </w:p>
  </w:footnote>
</w:footnotes>"""

        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
        footnotes_part = XmlPart(partname, content_type, None, package)
        footnotes_part._element = parse_xml(xml_content)

        part.package.main_document_part.relate_to(footnotes_part, footnote_rel_type)
        logger.info(f"[MANUAL] Created footnotes part")

    # Increment global counter for unique ID
    _footnote_counter += 1
    footnote_id = _footnote_counter

    # Get footnotes XML element
    if hasattr(footnotes_part, 'element'):
        fp_element = footnotes_part.element
    elif hasattr(footnotes_part, '_element'):
        fp_element = footnotes_part._element
    else:
        raise AttributeError("Cannot access footnotes element")

    # Create footnote content in footnotes.xml
    footnote = OxmlElement('w:footnote')
    footnote.set(qn('w:id'), str(footnote_id))

    # Footnote paragraph
    fp = OxmlElement('w:p')

    # Paragraph properties with Footnote Text style
    fpr = OxmlElement('w:pPr')
    fpstyle = OxmlElement('w:pStyle')
    fpstyle.set(qn('w:val'), 'FootnoteText')
    fpr.append(fpstyle)
    fp.append(fpr)

    # Footnote reference mark (the number at start of footnote)
    frun = OxmlElement('w:r')
    frun_ref = OxmlElement('w:rPr')
    rstyle = OxmlElement('w:rStyle')
    rstyle.set(qn('w:val'), 'FootnoteReference')
    frun_ref.append(rstyle)
    frun.append(frun_ref)
    frun.append(OxmlElement('w:footnoteRef'))
    fp.append(frun)

    # Space after number
    srun = OxmlElement('w:r')
    stext = OxmlElement('w:t')
    stext.set(qn('xml:space'), 'preserve')
    stext.text = ' '
    srun.append(stext)
    fp.append(srun)

    # Footnote text content
    ftrun = OxmlElement('w:r')
    fttext = OxmlElement('w:t')
    fttext.set(qn('xml:space'), 'preserve')
    fttext.text = text
    ftrun.append(fttext)
    fp.append(ftrun)

    footnote.append(fp)
    fp_element.append(footnote)

    # Add footnote reference marker in main text
    run = paragraph.add_run()
    r = run._r

    # Build w:rPr
    rPr = OxmlElement('w:rPr')

    # Footnote Reference style
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'FootnoteReference')
    rPr.append(rStyle)

    # Superscript
    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), 'superscript')
    rPr.append(vertAlign)

    # Font: Times New Roman
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)

    r.append(rPr)

    # Add footnoteReference with unique ID
    fref = OxmlElement('w:footnoteReference')
    fref.set(qn('w:id'), str(footnote_id))
    r.append(fref)

    logger.info(f"✅ [MANUAL SUCCESS] Footnote ID={footnote_id} added for citation {citation_id}")




def _generate_chart_comparison_content(doc: Document, report: Dict[str, Any]) -> None:
    # Title
    doc.add_heading(report.get('title', 'Comparative Analysis'), level=1)

    # Executive Summary
    if report.get('executive_summary'):
        _process_text(doc, report['executive_summary'])

    doc.add_page_break()

    # Iterate tasks
    for task in report.get('tasks', []):
        task_type = task.get('type')
        logger.info(f"Processing task: {task.get('id')} ({task_type})")

        if task_type == 'chart':
            if MATPLOTLIB_AVAILABLE:
                _add_chart_task(doc, task)
            else:
                doc.add_paragraph("[Chart generation unavailable - Matplotlib missing]")

        elif task_type == 'comparison':
            _add_comparison_task(doc, task)

        doc.add_paragraph() # Spacing

    # Bibliography
    if report.get('bibliography') or report.get('jurisprudence'):
        _add_bibliography(doc, report.get('bibliography', {}))

def _add_chart_task(doc: Document, task: Dict[str, Any]):
    label = task.get('label', 'Chart Task')
    doc.add_heading(label, level=2)

    spec = task.get('chart_spec', {})

    # Draw chart
    try:
        image_stream = _render_chart_from_spec(spec)
        if image_stream:
            doc.add_picture(image_stream, width=Cm(16))

            caption = spec.get('caption', f"Figure: {label}")
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p.style = 'Caption'
            except KeyError:
                pass # Fallback if style doesn't exist

    except Exception as e:
        logger.error(f"Failed to render chart: {e}")
        doc.add_paragraph(f"[Error rendering chart: {e}]")

def _render_chart_from_spec(spec: Dict[str, Any]) -> io.BytesIO:
    # Use matplotlib
    data = spec.get('data', {})
    chart_type = spec.get('chart_type', 'bar')
    options = spec.get('options', {})

    plt.figure(figsize=(10, 6))

    columns = data.get('columns', [])
    rows = data.get('rows', [])

    if not columns or not rows:
        return None

    # Assume first column is X, rest are Y series
    # rows is list of lists
    x_values = [str(row[0]) for row in rows]

    # Determine numeric columns keys
    y_labels = columns[1:]

    # Create numeric data structure
    # row = ["B1", 0.82, 0.78] -> x="B1", y1=0.82, y2=0.78

    # Basic bar chart logic
    if chart_type == 'bar':
        # If multiple series, we need offsets
        import numpy as np
        x_pos = np.arange(len(x_values))
        width = 0.35
        if len(y_labels) > 1:
            width = 0.8 / len(y_labels)

        for i, label in enumerate(y_labels):
            y_values = []
            for row in rows:
                try:
                    val = float(row[i + 1])
                except (ValueError, IndexError):
                    val = 0
                y_values.append(val)

            offset = (i - len(y_labels)/2 + 0.5) * width
            plt.bar(x_pos + offset, y_values, width, label=label)

        plt.xticks(x_pos, x_values, rotation=45, ha='right')

    elif chart_type == 'pie':
        # Pie chart usually 1 series. Take the first numeric column.
        y_values = []
        for row in rows:
            try:
                val = float(row[1])
            except (ValueError, IndexError):
                val = 0
            y_values.append(val)

        plt.pie(y_values, labels=x_values, autopct='%1.1f%%')

    else:
        # Default Line/Scatter
        for i, label in enumerate(y_labels):
            y_values = []
            for row in rows:
                try:
                    val = float(row[i + 1])
                except (ValueError, IndexError):
                    val = 0
                y_values.append(val)

            if chart_type == 'scatter':
                plt.scatter(x_values, y_values, label=label)
            else:
                plt.plot(x_values, y_values, marker='o', label=label)

        plt.xticks(rotation=45, ha='right')

    if options.get('title'):
        plt.title(options['title'])
    if options.get('x_label'):
        plt.xlabel(options['x_label'])
    if options.get('y_label'):
        plt.ylabel(options['y_label'])
    if options.get('legend', True) and chart_type != 'pie':
        plt.legend()

    plt.tight_layout()

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=300)
    plt.close()
    buf.seek(0)
    return buf

def _add_comparison_task(doc: Document, task: Dict[str, Any]):
    label = task.get('label', 'Comparison Task')
    doc.add_heading(label, level=2)

    # Table provided?
    comp_table = task.get('comparison_table', {})
    if comp_table:
        columns = comp_table.get('columns', [])
        rows = comp_table.get('rows', [])

        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Table Grid'

        # Header
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr_cells[i].text = str(col)
            # Make bold
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.bold = True

        # Rows
        for row in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                if i < len(columns):
                    row_cells[i].text = str(val)

        p = doc.add_paragraph(f"Table: {label}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.style = 'Caption'
        except KeyError:
            pass

    # Narrative
    if task.get('narrative_summary'):
        doc.add_heading("Analysis", level=3)
        _process_text(doc, task['narrative_summary'])

    if task.get('recommendation'):
        doc.add_heading("Recommendation", level=3)
        _process_text(doc, task['recommendation'])
